"""Mock end-to-end pipeline test.

Runs resume parsing, JD analysis, Anthropic-powered tailoring, and tracker
logging against a hardcoded sample JD. Does NOT touch a browser or submit
anything anywhere. Logs the result with status='test_dry_run'.
"""
import json
import os
import sys
from datetime import datetime
from pathlib import Path

# Force UTF-8 stdout on Windows so emoji and rich output don't crash cp1252.
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
        sys.stderr.reconfigure(encoding="utf-8")
    except Exception:
        pass

# Make project-root modules importable when running from src/
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from dotenv import load_dotenv
load_dotenv(ROOT / ".env")

import yaml
from rich.console import Console
from rich.panel import Panel

from utils import ensure_dirs
from resume_parser import parse_resume
from jd_analyzer import analyze_jd
from resume_tailor import tailor_resume, save_tailored_resume
from tracker import load_tracker, save_tracker, add_application


console = Console()


SAMPLE_JD = """\
Microsoft - Senior Cloud Solution Architect, AI Platform
Location: Remote (US)

About the Role
Microsoft is hiring a Senior Cloud Solution Architect focused on AI platform
adoption across our enterprise customer base. You will partner with account
teams to design and deliver Azure AI Foundry, AgentCore, and Copilot Studio
solutions, drive specialization attainment, and unblock customer production
deployments.

Responsibilities
- Lead technical pre-sales engagements for Azure AI Foundry and Agent platforms
- Design reference architectures for retrieval-augmented generation, agentic
  workflows, multi-agent orchestration, and responsible AI guardrails
- Drive co-sell motions with ISV partners
- Mentor field engineers on agent evaluation, prompt engineering, and LLMOps
- Contribute to specialization attainment and partner certification programs

Required Qualifications
- 7+ years in cloud architecture, solutions engineering, or technical sales
- Hands-on experience with Azure (AKS, App Service, Functions, AI Foundry)
- Production experience with at least one LLM provider (Azure OpenAI, Anthropic,
  AWS Bedrock)
- Strong Python skills; experience designing agentic systems
- Experience with vector databases and RAG patterns

Preferred Qualifications
- Experience with AWS AgentCore, Bedrock Agents, or LangGraph
- Background in customer-facing technical leadership
- Microsoft, AWS, or Google Cloud certifications
- Familiarity with MLOps and evaluation harnesses for LLM applications
"""


def _write_test_resume(path: Path) -> None:
    """Create a minimal but realistic test resume DOCX if no resume exists."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Jarrett Driscoll", level=0)
    doc.add_paragraph("jdriscollpro@gmail.com | 412-616-2093 | "
                       "linkedin.com/in/jarrettdriscoll | North Huntingdon, PA")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "Senior cloud solutions engineer with 8+ years building production "
        "systems on Azure and AWS. Recent focus on agentic AI: shipped LLM-"
        "powered automation pipelines, RAG retrieval services, and "
        "multi-agent orchestration for enterprise customers."
    )

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(
        "Python, TypeScript, Azure (AKS, Functions, App Service), AWS "
        "(Lambda, Bedrock), Anthropic Claude, OpenAI, LangChain, LangGraph, "
        "RAG, vector databases (pgvector, Pinecone), Docker, Kubernetes, "
        "Terraform, GitHub Actions, Playwright, FastAPI"
    )

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Senior Cloud Engineer — Acme Cloud Co. (2022–Present)", style="Heading 2")
    for b in [
        "Designed and shipped a multi-agent customer-support automation on "
        "Azure that handles 40k tickets/month with 92% deflection",
        "Built a RAG service over 2M enterprise docs using pgvector and "
        "Azure OpenAI, cutting analyst lookup time by 60%",
        "Led pre-sales architecture sessions with 30+ enterprise accounts",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("Cloud Solutions Engineer — Northbeam Systems (2018–2022)", style="Heading 2")
    for b in [
        "Migrated 15 legacy services to AKS with zero customer-visible downtime",
        "Built CI/CD on GitHub Actions used by 4 product teams",
        "Mentored 6 junior engineers on Azure architecture patterns",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.S. Computer Science, University of Pittsburgh, 2017")

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("Microsoft Certified: Azure Solutions Architect Expert")
    doc.add_paragraph("AWS Certified Solutions Architect — Associate")

    doc.save(path)


def _find_resume(data_dir: Path) -> Path:
    for name in ("resume_base.pdf", "resume_base.docx", "resume.pdf", "resume.docx"):
        p = data_dir / name
        if p.exists():
            return p

    test_resume = data_dir / "test_resume.docx"
    if not test_resume.exists():
        console.print(f"[yellow]No resume found in data/. Creating {test_resume.name}…[/yellow]")
        data_dir.mkdir(parents=True, exist_ok=True)
        _write_test_resume(test_resume)
    return test_resume


def _stub_analyze(jd_text: str, resume_data: dict) -> dict:
    """Deterministic offline analyzer for when no API key is available."""
    import re as _re
    jd_lower = jd_text.lower()
    resume_blob = " ".join([
        resume_data.get("summary") or "",
        " ".join(resume_data.get("skills") or []),
        " ".join(resume_data.get("experience") or []),
    ]).lower()

    candidate_terms = [
        "azure", "ai foundry", "agentcore", "copilot studio", "rag", "llm",
        "python", "aks", "anthropic", "bedrock", "langgraph", "co-sell",
        "specialization", "multi-agent", "vector", "responsible ai",
        "kubernetes", "openai", "agent",
    ]
    keywords = [t for t in candidate_terms if t in jd_lower]
    gaps = [t for t in keywords if t not in resume_blob]
    hits = [t for t in keywords if t in resume_blob]
    score = int(60 + 35 * (len(hits) / max(len(keywords), 1)))

    role_match = _re.search(r"(Senior [^\n,]+|Cloud Solution Architect[^\n]*)", jd_text)
    return {
        "company": "Microsoft",
        "role": role_match.group(0).strip() if role_match else "Senior Cloud Solution Architect",
        "required_skills": ["Azure", "Python", "LLM provider", "RAG", "Agentic systems"],
        "preferred_skills": ["AWS AgentCore", "LangGraph", "MLOps", "Microsoft certifications"],
        "keywords": keywords,
        "match_score": min(score, 95),
        "gaps": gaps,
        "summary": "Senior Cloud Solution Architect role at Microsoft focused "
                   "on Azure AI Foundry, AgentCore, and Copilot Studio "
                   "adoption across enterprise customers. Combines technical "
                   "pre-sales, reference architecture, and partner co-sell.",
    }


def _stub_tailor(resume_data: dict, jd_analysis: dict) -> str:
    """Offline tailor: emits markdown without calling the API."""
    name = resume_data.get("name") or "Candidate"
    contact = resume_data.get("contact", {})
    contact_line = " | ".join([v for v in (contact.get("email"),
                                              contact.get("phone"),
                                              contact.get("linkedin")) if v])
    skills = resume_data.get("skills") or []
    keywords = jd_analysis.get("keywords") or []
    # Reorder skills: keyword-matching first
    skills_sorted = sorted(skills, key=lambda s: 0 if any(k in s.lower() for k in keywords) else 1)

    parts = [f"# {name}"]
    if contact_line:
        parts.append(contact_line)
    parts += [
        "",
        "## Summary",
        f"Targeting: {jd_analysis.get('role')} at {jd_analysis.get('company')}.",
        resume_data.get("summary") or "",
        "",
        "## Skills",
        ", ".join(skills_sorted),
        "",
        "## Experience",
    ]
    parts += [f"- {line}" for line in (resume_data.get("experience") or [])]
    parts += ["", "## Education"]
    parts += [f"- {line}" for line in (resume_data.get("education") or [])]
    parts += ["", "## Certifications"]
    parts += [f"- {line}" for line in (resume_data.get("certifications") or [])]
    return "\n".join(parts)


def main() -> int:
    config_path = ROOT / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        config = yaml.safe_load(f)

    data_dir = ROOT / "data"
    tailored_dir = ROOT / config.get("resume", {}).get("tailored_output_dir", "data/tailored")
    tracker_path = ROOT / config.get("tracker", {}).get("path", "data/applications.json")

    ensure_dirs(str(tailored_dir), str(ROOT / "output" / "logs"))

    console.print(Panel.fit("[bold cyan]Mock Pipeline Test — DRY RUN[/bold cyan]\n"
                              "No browser, no submission, no real URL fetch."))

    resume_path = _find_resume(data_dir)
    console.print(f"[cyan]📄 Parsing resume:[/cyan] {resume_path}")
    resume_data = parse_resume(str(resume_path))
    console.print(f"   name: {resume_data.get('name')}")
    console.print(f"   skills detected: {len(resume_data.get('skills') or [])}")
    console.print(f"   experience lines: {len(resume_data.get('experience') or [])}")

    model = config.get("anthropic", {}).get("model", "claude-sonnet-4-20250514")
    use_stub = not os.environ.get("ANTHROPIC_API_KEY")

    if use_stub:
        console.print(
            "\n[yellow]⚠️  ANTHROPIC_API_KEY not set — using deterministic "
            "stub analyzer/tailor so the pipeline can still be exercised "
            "end-to-end. Set the key and re-run for real LLM output.[/yellow]"
        )

    console.print(f"\n[cyan]🔍 Analyzing sample JD (model={model})…[/cyan]")
    if use_stub:
        jd_analysis = _stub_analyze(SAMPLE_JD, resume_data)
    else:
        jd_analysis = analyze_jd(SAMPLE_JD, resume_data, model)
    console.print("\n[bold]JD Analysis JSON:[/bold]")
    console.print_json(data=jd_analysis)

    console.print(Panel(
        f"[bold]{jd_analysis.get('company')} - {jd_analysis.get('role')}[/bold]\n"
        f"Match Score: [green]{jd_analysis.get('match_score')}%[/green]\n"
        f"Gaps: {', '.join(jd_analysis.get('gaps', [])[:8])}\n"
        f"Keywords: {', '.join(jd_analysis.get('keywords', [])[:10])}\n\n"
        f"{jd_analysis.get('summary', '')}",
        title="Match Summary",
        border_style="cyan",
    ))

    console.print(f"\n[cyan]✏️  Tailoring resume…[/cyan]")
    if use_stub:
        tailored_md = _stub_tailor(resume_data, jd_analysis)
    else:
        tailored_md = tailor_resume(resume_data, jd_analysis, model)
    tailored_path = save_tailored_resume(
        tailored_md,
        jd_analysis.get("company", "company"),
        jd_analysis.get("role", "role"),
        str(tailored_dir),
    )
    md_path = tailored_path.replace(".docx", ".md")
    console.print(f"[green]✅ Tailored resume saved: {tailored_path}[/green]")
    console.print(f"   markdown copy: {md_path}")

    fake_url = "https://test.local/dry-run/" + datetime.now().strftime("%Y%m%dT%H%M%S")
    tracker = load_tracker(str(tracker_path))
    tracker = add_application(
        tracker,
        company=jd_analysis.get("company", "Unknown"),
        role=jd_analysis.get("role", "Unknown"),
        url=fake_url,
        match_score=int(jd_analysis.get("match_score", 0)),
        keywords_added=jd_analysis.get("keywords", []),
        tailored_resume=tailored_path,
        status="test_dry_run",
        notes="Generated by src/test_pipeline.py — no browser, no submission.",
    )
    save_tracker(tracker, str(tracker_path))
    new_entry = tracker["applications"][-1]

    console.print(f"\n[bold]Tracker entry written to {tracker_path}:[/bold]")
    console.print_json(data=new_entry)

    console.print(Panel.fit(
        "[bold green]Pipeline Summary[/bold green]\n"
        f"• Resume parsed: {resume_path}\n"
        f"• JD source: hardcoded sample ({len(SAMPLE_JD)} chars)\n"
        f"• Match score: {jd_analysis.get('match_score')}%\n"
        f"• Keyword gaps: {', '.join(jd_analysis.get('gaps', []) or ['none'])}\n"
        f"• Tailored resume: {tailored_path}\n"
        f"• Tracker entry id: {new_entry['id']}\n"
        f"• Status: test_dry_run (no browser launched, nothing submitted)",
        border_style="green",
    ))
    return 0


if __name__ == "__main__":
    sys.exit(main())
