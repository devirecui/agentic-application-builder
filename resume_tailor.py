import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
from anthropic import Anthropic
from utils import slugify, log_token_usage

load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))


TAILOR_PROMPT = """You are tailoring a resume for a specific job application.

ORIGINAL RESUME:
{resume_json}

JD ANALYSIS:
{jd_analysis}

Rewrite the resume to:
1. Emphasize experience most relevant to this role
2. Naturally incorporate these keywords: {keywords}
3. Reorder skills to put most relevant first
4. Strengthen bullet points with the role's language
5. Never fabricate experience or credentials

Return the tailored resume as clean markdown with sections:
# {{name}}
contact line
## Summary
## Skills
## Experience
## Education
## Certifications

Do not add skills or experience that don't exist in the original."""


def tailor_resume(resume_data: dict, jd_analysis: dict, model: str = "claude-sonnet-4-6") -> str:
    client = Anthropic()

    safe_resume = {k: v for k, v in resume_data.items() if k != "raw_text"}
    prompt = TAILOR_PROMPT.format(
        resume_json=json.dumps(safe_resume, indent=2),
        jd_analysis=json.dumps(jd_analysis, indent=2),
        keywords=", ".join(jd_analysis.get("keywords", [])),
    )

    last_err = None
    for _ in range(2):
        try:
            msg = client.messages.create(
                model=model,
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}],
            )
            log_token_usage("resume_tailor", model, msg.usage.input_tokens, msg.usage.output_tokens)
            return "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()
        except Exception as e:
            last_err = e
            time.sleep(1)
    raise RuntimeError(f"Anthropic API failed: {last_err}")


def save_tailored_resume(content: str, company: str, role: str, output_dir: str) -> str:
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d")
    fname = f"{slugify(company)}_{slugify(role)}_{stamp}.docx"
    path = os.path.join(output_dir, fname)
    _markdown_to_docx(content, path)

    md_path = path.replace(".docx", ".md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    return path


def _markdown_to_docx(md: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^[-*]\s", line):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(out_path)
